import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="Smart AI Financial Advisor", layout="wide")
st.title("📈 Smart AI Financial and Crypto Advisor ")

# ------------------------------
# 📁 Upload Portfolio
# ------------------------------
uploaded_file = st.file_uploader("Upload your crypto portfolio CSV", type=["csv"])

if uploaded_file:
    df = pd.read_csv(uploaded_file)

    # Simulated live prices
    price_lookup = {
        "bitcoin": 35000,
        "ethereum": 2100,
        "ripple": 0.6,
        "cardano": 0.35,
        "solana": 22
    }

    df["live_price"] = df["asset"].str.lower().map(price_lookup)
    df["current_value"] = df["quantity"] * df["live_price"]
    df["invested"] = df["quantity"] * df["avg_buy_price"]
    df["gain"] = df["current_value"] - df["invested"]
    df["gain_percent"] = (df["gain"] / df["invested"]) * 100

    st.subheader("📊 Portfolio Summary")
    st.dataframe(df)

    # ------------------------------
    # 📈 Visual Charts
    # ------------------------------
    st.subheader("📊 Visual Charts")

    pie_buf = BytesIO()
    df.set_index("asset")["current_value"].plot.pie(autopct="%1.1f%%", figsize=(5, 5))
    plt.title("Portfolio Allocation")
    plt.ylabel("")
    plt.tight_layout()
    plt.savefig(pie_buf, format="png")
    st.image(pie_buf)

    bar_buf = BytesIO()
    df.plot.bar(x="asset", y="gain", color="green", legend=False)
    plt.title("Gain/Loss by Asset")
    plt.ylabel("Gain ($)")
    plt.tight_layout()
    plt.savefig(bar_buf, format="png")
    st.image(bar_buf)

    # ------------------------------
    # 🧠 Insights
    # ------------------------------
    st.subheader("🧠 Portfolio Insights")

    def generate_insights(data):
        total_gain = data["gain"].sum()
        avg_gain = data["gain_percent"].mean()
        dominant_asset = data.loc[data["current_value"].idxmax()]["asset"]
        dominant_value = data["current_value"].max()
        portfolio_value = data["current_value"].sum()
        concentration = (dominant_value / portfolio_value) * 100

        insights = []

        if total_gain > 0:
            insights.append(f"Your portfolio is in profit with a total gain of ${total_gain:,.2f}.")
        elif total_gain < 0:
            insights.append(f"Your portfolio is running at a loss of ${-total_gain:,.2f}.")

        if avg_gain > 15:
            insights.append("Average gains per asset are high — consider booking some profits.")
        elif avg_gain < -10:
            insights.append("Average losses suggest potential reevaluation of your holdings.")

        if concentration > 60:
            insights.append(f"Over {concentration:.1f}% of your portfolio is in {dominant_asset}. Consider rebalancing.")

        if len(data) < 3:
            insights.append("Portfolio is under-diversified. Add more assets to reduce risk.")

        if not insights:
            insights.append("Portfolio looks balanced and healthy based on current data.")

        return "\n".join(insights)

    advice = generate_insights(df)
    st.write(advice)

    # ------------------------------
    # 📄 Word Report Generator
    # ------------------------------
    st.subheader("📥 Download Word Report")
    if st.button("Generate Report"):
        doc = Document()
        doc.add_heading("Crypto Portfolio Analysis Report", 0)

        doc.add_heading("Portfolio Table", level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = f"{item:.2f}" if isinstance(item, float) else str(item)

        doc.add_heading("Charts", level=1)
        pie_buf.seek(0)
        bar_buf.seek(0)
        doc.add_picture(pie_buf, width=Inches(5))
        doc.add_picture(bar_buf, width=Inches(5))

        doc.add_heading("Insights", level=1)
        doc.add_paragraph(advice)

        word_buf = BytesIO()
        doc.save(word_buf)
        word_buf.seek(0)

        st.download_button(
            label="📄 Download Word Report",
            data=word_buf,
            file_name="portfolio_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
